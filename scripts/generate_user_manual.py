#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成《网优百宝箱》操作说明 Word 文档"""

from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
except ImportError:
    raise SystemExit("请先安装: pip install python-docx")

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "网优百宝箱_操作说明.docx"


def set_doc_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "PingFang SC"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(11)


def add_title(doc, text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def add_h(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_p(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()


def build():
    doc = Document()
    set_doc_font(doc)

    add_title(doc, "网优百宝箱 操作说明")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"版本 v1.2.1　　文档日期 {date.today().isoformat()}")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x66, 0x77, 0x88)
    doc.add_page_break()

    add_h(doc, "1. 产品简介")
    add_p(
        doc,
        "网优百宝箱是一款轻量化、可离线部署的 4G/5G 小区规划工具，面向无线网络优化工程师。"
        "主要能力包括：工参与网管数据管理、GIS 地图可视化、三级 PCI 智能规划、7:3 加权邻区规划、"
        "多制式邻区生成（4G↔4G、4G↔5G 等）、冲突与干扰校验、规划结果与 MML 脚本导出，"
        "并适配陆地及近海超远覆盖等场景。",
    )
    add_bullets(
        doc,
        [
            "功能分析：小区规划（PCI+邻区）、PCI 干扰分析",
            "数据导入：工参导入、网管导入配置、网管数据查看",
            "技术特点：浏览器访问、单端口部署、无需前端构建",
        ],
    )

    add_h(doc, "2. 运行环境与安装")
    add_h(doc, "2.1 环境要求", level=2)
    add_bullets(
        doc,
        [
            "操作系统：Windows / macOS / Linux",
            "Python：3.9 及以上",
            "浏览器：Chrome、Edge、Safari 等现代浏览器",
            "网络：地图底图需联网；核心业务可本地离线运行",
        ],
    )
    add_h(doc, "2.2 安装依赖", level=2)
    add_p(doc, "在项目根目录打开终端，执行：")
    add_p(doc, "pip install -r requirements.txt", bold=True)
    add_h(doc, "2.3 启动与停止", level=2)
    add_table(
        doc,
        ["操作", "命令", "说明"],
        [
            ["启动（前台）", "bash start.sh", "默认端口 4001，绑定 0.0.0.0"],
            ["停止", "bash stop.sh", "停止当前服务进程"],
            ["重启（后台）", "bash restart.sh", "日志输出至 logs/app.out"],
            ["自定义端口", "PORT=8888 bash start.sh", "可同时设置 HOST=0.0.0.0"],
        ],
    )
    add_p(doc, "启动成功后，在浏览器访问：http://localhost:4001（或您设置的端口）。")
    add_p(doc, "API 接口文档：http://localhost:4001/docs")

    add_h(doc, "3. 首页与功能导航")
    add_p(doc, "打开系统首页后，通过功能卡片进入各模块：")
    add_table(
        doc,
        ["模块", "入口", "用途"],
        [
            ["小区规划", "功能分析 → 小区规划", "PCI + 邻区规划、单站/批量规划、地图展示"],
            ["PCI 干扰分析", "功能分析 → PCI 干扰分析", "干扰检测、全网/局部 PCI 重规划"],
            ["工参导入", "数据导入 → 工参导入", "4G/5G 工参上传、模板、导出"],
            ["网管导入配置", "数据导入 → 网管导入配置", "网管 Excel 多 Sheet 导入与列映射"],
            ["网管数据查看", "数据导入 → 网管数据查看", "已导入表浏览、筛选、导出"],
        ],
    )

    add_h(doc, "4. 工参导入（使用前必读）")
    add_p(doc, "进行小区规划或 PCI 分析前，需先导入工参数据。")
    add_h(doc, "4.1 推荐流程", level=2)
    add_numbered(
        doc,
        [
            "进入「工参导入」页面。",
            "下载模板：可按 4G、5G 或双制式下载 Excel 模板。",
            "按模板填写小区信息（必填项见模板「字段说明」）。",
            "选择「追加」或「替换」模式后上传文件。",
            "在列表中确认导入条数与字段是否正确，必要时导出核对。",
        ],
    )
    add_h(doc, "4.2 主要工参字段", level=2)
    add_table(
        doc,
        ["字段", "说明", "示例"],
        [
            ["ECGI", "小区全局标识", "460-00-12345-1"],
            ["小区名称", "站点名-扇区", "BJ001-1"],
            ["制式", "LTE / NR", "LTE"],
            ["频点", "频段或中心频率", "1850"],
            ["经度 / 纬度", "WGS84", "116.404 / 39.915"],
            ["方位角", "正北顺时针 0–360°", "60"],
            ["覆盖半径", "米", "500"],
            ["PCI", "物理小区 ID", "100"],
            ["站点类型", "陆地 / 近海 / 室内等", "陆地"],
        ],
    )
    add_h(doc, "4.3 扩展字段（单站/批量规划）", level=2)
    add_table(
        doc,
        ["字段", "必填", "说明"],
        [
            ["扇区数", "否", "1–6，默认 1"],
            ["基方位角", "否", "0–360，默认 0"],
            ["规划类型", "否", "宏站 / 微站 / 室分"],
            ["邻区规划", "否", "多选用 | 分隔：4G_4G|4G_5G|5G_4G|5G_5G"],
            ["锁定", "否", "「是」则该小区不参与 PCI 重分配"],
        ],
    )

    add_h(doc, "5. 小区规划（PCI + 邻区）")
    add_p(doc, "路径：首页 → 小区规划。页面为地图大屏 + 右侧详情与操作日志。")
    add_h(doc, "5.1 地图与图层", level=2)
    add_bullets(
        doc,
        [
            "扇区按 PCI mod3 着色（红/黄/蓝）；冲突小区标红闪烁。",
            "可开启 4G/5G PCI 数字标签图层，便于核对。",
            "邻区连线按类型分色：4G↔4G、4G↔5G、5G↔4G、5G↔5G。",
            "地图下方「邻区过滤」可按类型与得分阈值筛选显示的邻区连线。",
            "工具栏：重置视图、聚焦 4G、聚焦 5G。",
        ],
    )
    add_h(doc, "5.2 单站规划", level=2)
    add_numbered(
        doc,
        [
            "点击「单站规划」打开浮动面板（可拖动、缩放）。",
            "选择规划模式：PCI+邻区 / 仅 PCI / 仅邻区。",
            "填写制式、频段、站点类型（宏站/微站/室分）、经纬度、扇区数、方位角。",
            "勾选需要的邻区类型（4G↔4G 等）及得分阈值。",
            "点击「规划」：地图聚焦约 10 km 范围展示规划结果。",
            "点击「导出 xlsx」：下载 PCI 规划表及按邻区类型分 Sheet 的结果。",
        ],
    )
    add_h(doc, "5.3 批量规划", level=2)
    add_numbered(
        doc,
        [
            "点击「批量规划」，下载 4G/5G/双制式模板之一。",
            "按表头标注的 [必填]/[可选]/[枚举] 填写（单次建议不超过 500 行）。",
            "选择 Excel/CSV 文件，点击「规划并导出」。",
            "浏览器将直接下载包含多 Sheet 的结果文件。",
        ],
    )
    add_h(doc, "5.4 邻区参数", level=2)
    add_table(
        doc,
        ["参数", "默认值", "说明"],
        [
            ["最大邻区数", "16", "每小区保留的邻区数量上限"],
            ["最大距离(km)", "5.0", "邻区搜索半径"],
            ["距离权重", "0.7", "与交叠权重之和建议为 1"],
            ["交叠权重", "0.3", "扇区覆盖交叠占比权重"],
            ["得分阈值", "0.10", "低于阈值的邻区关系可被过滤"],
            ["异系统邻区", "开启", "是否规划 4G↔5G 等跨制式邻区"],
        ],
    )
    add_h(doc, "5.5 站点类型与 PCI 距离规则", level=2)
    add_table(
        doc,
        ["类型", "Mod3 安全距离", "同 PCI 最小距离", "典型场景"],
        [
            ["宏站 macro", "700 m", "5 km", "室外宏蜂窝"],
            ["微站 micro", "200 m", "3 km", "小微站"],
            ["室分 indoor", "100 m", "2 km", "室内分布"],
        ],
    )
    add_p(doc, "上述阈值与全局默认取更严格者，用于 PCI 分配与冲突判断。")

    add_h(doc, "6. PCI 干扰分析")
    add_p(doc, "路径：首页 → PCI 干扰分析。适合在已有工参基础上做干扰排查与 PCI 重规划。")
    add_h(doc, "6.1 干扰分析", level=2)
    add_numbered(
        doc,
        [
            "打开「干扰分析」浮动面板，设置分析距离、交叠阈值等。",
            "勾选检测项：同频、邻频、PCI 冲突、Mod3/Mod6 等。",
            "可选：框选矩形或圆形区域，仅分析选定范围。",
            "点击「分析干扰」，地图与列表面板展示 issues 与严重度。",
            "点击「导出 xlsx」生成干扰分析报告（含 mitigation 建议）。",
        ],
    )
    add_h(doc, "6.2 PCI 规划与冲突处理", level=2)
    add_numbered(
        doc,
        [
            "在「PCI 规划」面板选择算法引擎：legacy（默认）或 rftools。",
            "按需勾选 Mod6、Mod30（NR DMRS）等检查项。",
            "「全网规划」：对全部未锁定小区重新分配 PCI 并刷新地图。",
            "「局部微调」：先框选/圆选区域，仅对选区内小区重规划。",
            "规划完成后可进行「冲突校验」；若有冲突，查看冲突清单弹窗并按要求调整或再次规划。",
        ],
    )

    add_h(doc, "7. 网管数据导入与查看")
    add_h(doc, "7.1 网管导入配置", level=2)
    add_numbered(
        doc,
        [
            "进入「网管导入配置」。",
            "步骤 1：点击或拖拽上传网管 Excel（支持多文件队列）。",
            "步骤 2：勾选需要导入的 Sheet。",
            "步骤 3：点击「导入选中的 Sheet」。",
            "列映射：通过「列配置」从 YAML 预定义或从 Excel 读取表头，保存后供后续导入复用。",
            "右侧面板可查看「已导入表」与「导入历史」，支持刷新。",
        ],
    )
    add_h(doc, "7.2 网管数据查看", level=2)
    add_p(doc, "在「网管数据查看」中选择已导入的数据表，使用搜索与分页浏览记录，必要时导出为文件供外部分析。")

    add_h(doc, "8. 导出物说明")
    add_table(
        doc,
        ["导出类型", "常见入口", "内容"],
        [
            ["工参导出", "工参导入页", "当前库内工参 Excel"],
            ["分 Sheet 规划结果", "单站规划 → 导出 xlsx", "PCI 表 + 各邻区类型 Sheet"],
            ["批量规划结果", "批量规划 → 规划并导出", "多 Sheet 一站式下载"],
            ["干扰分析报告", "PCI 页 → 导出 xlsx", "干扰项、统计、优化建议"],
            ["冲突报表 / MML", "规划流程中的导出菜单", "冲突清单或网管脚本（视功能开放项）"],
        ],
    )

    add_h(doc, "9. 典型业务场景")
    add_h(doc, "9.1 新站入网 PCI + 邻区", level=2)
    add_numbered(
        doc,
        [
            "工参导入存量网络；",
            "小区规划 → 单站规划，输入新站经纬度与扇区；",
            "执行规划并导出，提交现场或网管配置。",
        ],
    )
    add_h(doc, "9.2 批量新开站", level=2)
    add_numbered(
        doc,
        [
            "下载双制式批量模板，填写多行新站；",
            "批量规划并导出；",
            "按 Sheet 分发给不同制式负责人。",
        ],
    )
    add_h(doc, "9.3 现网 PCI 干扰整治", level=2)
    add_numbered(
        doc,
        [
            "确认工参为最新；",
            "PCI 干扰分析 → 区域分析或全网分析；",
            "全网或局部 PCI 重规划 → 冲突校验 → 导出报告与工参。",
        ],
    )

    add_h(doc, "10. 常见问题")
    add_table(
        doc,
        ["现象", "可能原因", "处理建议"],
        [
            ["无法打开页面", "服务未启动或端口占用", "执行 start.sh，或更换 PORT"],
            ["地图空白", "底图服务不可达", "检查网络；确认浏览器未拦截混合内容"],
            ["上传失败", "表头与模板不一致", "重新下载模板，对照「字段说明」"],
            ["规划很慢", "小区数量大", "先用局部/单站验证参数，再全网规划"],
            ["邻区过少", "得分阈值过高或距离过小", "调低阈值或增大最大距离"],
            ["PCI 仍有冲突", "锁定小区过多或站型阈值严", "减少锁定，或扩大规划范围后重试"],
        ],
    )

    add_h(doc, "11. 附录")
    add_h(doc, "11.1 目录结构（简要）", level=2)
    add_p(doc, "backend/ 后端服务；frontend/ 前端页面；static/ 示例数据；config.yaml 全局配置；logs/ 运行日志。")
    add_h(doc, "11.2 联系与版本", level=2)
    add_p(doc, "软件名称：网优百宝箱　版本：v1.2.1")
    add_p(doc, "本文档随软件功能更新，若界面与文档不一致，以实际界面为准。")

    doc.save(OUT)
    print(f"已生成: {OUT}")


if __name__ == "__main__":
    build()